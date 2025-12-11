
import os
from docx import Document
from docx.shared import Inches

def export_word(kpis: dict, schedule_table, chart_path: str, outfile: str = "CHP_Report.docx"):
    doc = Document()
    doc.add_heading('CHP Feasibility & Financial Summary', level=1)

    doc.add_paragraph(f"Engine: {kpis.get('engine_display','')}")
    doc.add_paragraph(f"Rated Power: {kpis.get('rated_power_kw','')} kW")
    doc.add_paragraph(f"Electrical Efficiency (HHV): {kpis.get('elec_eff_pct','')}%")
    doc.add_paragraph(f"Total System Efficiency: {kpis.get('total_eff_pct','')}%")
    doc.add_paragraph(f"ITC Applied: {kpis.get('itc_pct','')}%")

    doc.add_heading('Key Financials', level=2)
    p = doc.add_paragraph()
    p.add_run(f"IRR: {kpis['irr']:.2%}\n")
    p.add_run(f"NPV @ {kpis['npv_rate']:.0%}: ${kpis['npv']:,.0f}\n")
    p.add_run(f"Simple Payback: {kpis['payback']} years\n")
    p.add_run(f"Discounted Payback: {kpis['disc_payback']} years\n")

    doc.add_heading('Annual Cash Flow', level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Year'; hdr_cells[1].text = 'Revenue ($)'
    hdr_cells[2].text = 'Fuel ($)'; hdr_cells[3].text = 'O&M ($)'
    hdr_cells[4].text = 'Debt ($)'; hdr_cells[5].text = 'Net CF ($)'

    for row in schedule_table:
        r = table.add_row().cells
        r[0].text = str(row['year'])
        r[1].text = f"{(row['rev_elec']+row['rev_therm']):,.0f}"
        r[2].text = f"{row['fuel_cost']:,.0f}"
        r[3].text = f"{row['om_cost']:,.0f}"
        r[4].text = f"{row['debt_service']:,.0f}"
        r[5].text = f"{row['net_cf']:,.0f}"

    if chart_path and os.path.exists(chart_path):
        doc.add_heading('Annual Cash Flow Chart', level=2)
        doc.add_picture(chart_path, width=Inches(6.0))

    doc.add_paragraph("Methodology aligns with EPA CHP efficiency framework (total system efficiency).")
    doc.save(outfile)
    return outfile
